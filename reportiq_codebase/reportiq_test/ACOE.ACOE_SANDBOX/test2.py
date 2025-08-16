from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown2
from bs4 import BeautifulSoup
import streamlit as st

def markdown_to_docx_with_tables_fixed(markdown_text):
    """Convert Markdown to DOCX, preserving headings, paragraphs, lists, and tables with proper nesting and formatting."""
    try:
        html = markdown2.markdown(markdown_text, extras=["tables", "fenced-code-blocks", "strike", "underline"])
        soup = BeautifulSoup(html, 'html.parser')
        doc = Document()

        list_number = 1  # For numbered lists

        def add_list_item(text, style, indent_level=0):
            p = doc.add_paragraph(style=style)
            run = p.add_run(text.lstrip())
            run.font.size = Pt(11)
            p.paragraph_format.left_indent = Inches(0.15 * indent_level)
            p.paragraph_format.first_line_indent = Inches(0.15 * indent_level)
            p.paragraph_format.space_after = Pt(2)

        def add_heading(text, level):
            p = doc.add_heading('', level=level)
            run = p.add_run(text.strip())
            run.bold = True
            p.paragraph_format.space_after = Pt(6)

        def add_paragraph(text):
            p = doc.add_paragraph(text.strip())
            p.paragraph_format.space_after = Pt(6)


        def process_ul(ul_element, indent_level=0):
            for li in ul_element.find_all('li', recursive=False):
                # Extract text excluding nested lists
                text_parts = [str(c) for c in li.contents if c.name not in ['ul', 'ol']]
                text = BeautifulSoup(''.join(text_parts), 'html.parser').get_text().strip()
                add_list_item(text, style='List Bullet', indent_level=indent_level)

                # Recursively process nested lists
                for nested in li.find_all(['ul', 'ol'], recursive=False):
                    if nested.name == 'ul':
                        process_ul(nested, indent_level + 1)
                    elif nested.name == 'ol':
                        process_ol(nested, indent_level + 1)

        def process_ol(ol_element, indent_level=0):
            nonlocal list_number
            for li in ol_element.find_all('li', recursive=False):
                text_parts = [str(c) for c in li.contents if c.name not in ['ul', 'ol']]
                text = BeautifulSoup(''.join(text_parts), 'html.parser').get_text().strip()
                add_list_item(f"{list_number}. {text}", style='Normal', indent_level=indent_level)
                list_number += 1

                for nested in li.find_all(['ul', 'ol'], recursive=False):
                    if nested.name == 'ul':
                        process_ul(nested, indent_level + 1)
                    elif nested.name == 'ol':
                        process_ol(nested, indent_level + 1)


        for element in soup.contents:
            if not hasattr(element, 'name'):
                continue

            if element.name in ['h1', 'h2', 'h3']:
                level = int(element.name[1])
                add_heading(element.text.strip(), level)
                list_number = 1  # Reset numbering at each heading

            elif element.name == 'p':
                add_paragraph(element.text.strip())

            elif element.name == 'ul':
                process_ul(element)

            elif element.name == 'ol':
                process_ol(element)

            elif element.name == 'table':
                rows = element.find_all('tr')
                if rows:
                    cols = rows[0].find_all(['td', 'th'])
                    table = doc.add_table(rows=1, cols=len(cols))
                    table.style = 'Table Grid'

                    hdr_cells = table.rows[0].cells
                    for i, cell in enumerate(cols):
                        hdr_cells[i].text = cell.text.strip()

                    for row in rows[1:]:
                        cells = row.find_all(['td', 'th'])
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(cells):
                            row_cells[i].text = cell.text.strip()

                    # Add spacing after table
                    doc.add_paragraph()

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as e:
        print(f"Error converting to DOCX: {e}")
        return None
markdown_text = """
# Power BI Report Documentation

## Overview
This Power BI report provides insights into lab turnaround times for emergency departments (EDs) across various locations, procedures, and time periods. The report is designed to help stakeholders analyze and compare median and mean turnaround times in minutes for lab results. It includes visualizations for trends, comparisons, and detailed breakdowns by state, hospital, and procedure.

### Key Features:
- **Median and Mean Turnaround Times**: Displays system-wide and location-specific metrics.
- **Trend Analysis**: Visualizes turnaround times over the last 24 weeks.
- **Comparative Analysis**: Highlights locations performing above or below system benchmarks.
- **Drill-Down Functionality**: Allows users to explore data at the procedure level.
- **Interactive Filters**: Enables filtering by state, hospital, procedure, and time period.

---

## Key Contacts
| Role                | Name                  | Contact Information                                   |
|---------------------|-----------------------|-----------------------------------------------------|
| Report Owner        | Analytics Center of Excellence (ACOE) | [Feedback for ACOE](https://providence4.sharepoint.com/sites/AnalyticsCenterofExcellenceACOE/SitePages/Feedback-for-the-ACOE.aspx) |
| Developer           | Not Specified         | Not Specified                                       |
| Data Source Owner   | ACOE_DEV Database Team| Not Specified                                       |

---

## Data Flow
1. **Data Source**: The data is sourced from the `LAB_ORDER_TURNAROUND` table in the `RPT_LAB` schema of the `ACOE_DEV` database hosted on Snowflake.
2. **Data Import**: Data is imported into Power BI using direct queries or scheduled refreshes.
3. **Transformations**:
   - Calculated columns such as `WeekDayName`, `WeekOrder`, and `States` are derived for better visualization.
   - Measures like `System Median`, `System Mean`, and `color` are calculated for comparative analysis.
4. **Visualization**:
   - The report includes bar charts, line charts, slicers, and cards for interactive data exploration.

---

## Data Explorer
### Pages in the Report:
1. **Lab Turnaround Median Time by Location**:
   - Displays median turnaround times by state, hospital, and procedure.
        - hello alllll
   - Includes a bar chart with drill-down capabilities.
        - hi all
2. **Lab Turnaround Mean Time by Location**:
   - Similar to the median page but focuses on mean turnaround times.
3. **Trend Analysis**:
   - Line charts showing 24-week trends for turnaround times.
4. **Tooltips**:
   - Provides detailed information on hover for specific data points.

### Filters:
- **State(s)**: Dropdown filter for selecting one or multiple states.
- **Hospital**: Dropdown filter for selecting specific hospitals.
- **Procedure**: Dropdown filter for selecting lab procedures.
- **Time Period**: Relative date slicer for filtering by calendar months.

---

## Data Dictionary

### Tables and Columns

#### Table: `VW_LAB_ORDER_TURNAROUND`
| Column Name              | Data Type   | Description                                                                                     |
|--------------------------|-------------|-------------------------------------------------------------------------------------------------|
| `SILO`                  | String      | Represents the test group.                                                                     |
| `COLLECTING_SERVICE_AREA` | String      | Service area where the lab order was collected.                                                |
| `COLLECTING_HOSPITAL`    | String      | Hospital where the lab order was collected.                                                    |
| `RESULTING_LAB`          | String      | Lab responsible for processing the result.                                                     |
| `PROC_NAME`              | String      | Name of the procedure performed.                                                              |
| `ORDER_ID`               | Double      | Unique identifier for each lab order.                                                         |
| `RESULT_DATE`            | DateTime    | Date when the result was finalized.                                                           |
| `RCV_TO_RSLT`            | Double      | Time (in minutes) from receiving the sample to finalizing the result.                         |
| `RESULT_DT`              | DateTime    | General date representation of the result date.                                               |
| `CREATE_TS`              | DateTime    | Timestamp indicating when the record was created.                                              |
| `Service Area`           | Calculated  | Cleaned version of `COLLECTING_SERVICE_AREA`.                                                  |
| `Hospital`               | Calculated  | Cleaned version of `COLLECTING_HOSPITAL` with standardized names.                              |
| `States`                 | Calculated  | Categorized states based on service areas (e.g., California, Oregon).                         |
| `WeekDayName`            | Calculated  | Day of the week derived from `RESULT_DATE`.                                                   |
| `WeekOrder`              | Calculated  | Numerical representation of weekdays (Sunday = 0, Monday = 1, etc.).                          |
| `WeekID`                 | Calculated  | Week number derived from `RESULT_DT`.                                                         |
| `MonthWeek`              | Calculated  | Concatenation of month and year derived from `RESULT_DT`.                                      |

#### Measures
| Measure Name             | Description                                                                                     |
|--------------------------|-------------------------------------------------------------------------------------------------|
| `test`                   | Returns the median of `RCV_TO_RSLT` if there is only one value; otherwise returns blank.        |
| `System Median`          | Median turnaround time across all procedures.                                                  |
| `System Mean`            | Mean turnaround time across all procedures.                                                    |
| `color`                  | Binary indicator (1/0) based on whether a location's median exceeds the system median.          |
| `color mean`             | Binary indicator (1/0) based on whether a location's mean exceeds the system mean.              |

---

This documentation provides a comprehensive overview of the report, its structure, and its data components to ensure clarity and usability for stakeholders.
"""

# Streamlit download button
docx_file = markdown_to_docx_with_tables_fixed(markdown_text)
if docx_file:
    st.download_button(
        label="Download DOCX",
        data=docx_file,
        file_name="documentation.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

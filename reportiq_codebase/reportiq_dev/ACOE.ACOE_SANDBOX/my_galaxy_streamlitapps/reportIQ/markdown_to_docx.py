from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
import markdown2
import streamlit as st
from bs4 import BeautifulSoup
# add this to requirements.txt : pip install python-docx markdown2 beautifulsoup4
def markdown_to_docx(markdown_text):
    """Convert Markdown to DOCX, preserving headings, paragraphs, lists, and tables with proper nesting and formatting."""
    try:
        html = markdown2.markdown(markdown_text, extras=["tables", "fenced-code-blocks", "strike", "underline"])
        soup = BeautifulSoup(html, 'html.parser')
        doc = Document()
        list_number = 1

        def add_list_item(text, style, indent_level):
            p = doc.add_paragraph(style=style)
            run = p.add_run(text.lstrip())
            run.font.size = Pt(11)
            indent = Inches(0.15 * indent_level)
            p.paragraph_format.left_indent = indent
            p.paragraph_format.first_line_indent = indent
            p.paragraph_format.space_after = Pt(2)

        def process_list(element, indent_level, ordered=False):
            nonlocal list_number
            for li in element.find_all('li', recursive=False):
                text = BeautifulSoup(''.join(str(c) for c in li.contents if c.name not in ['ul', 'ol']), 'html.parser').get_text().strip()
                if ordered:
                    add_list_item(f"{list_number}. {text}", style='Normal', indent_level=indent_level)
                    list_number += 1
                else:
                    add_list_item(text, style='List Bullet', indent_level=indent_level)

                for nested in li.find_all(['ul', 'ol'], recursive=False):
                    process_list(nested, indent_level + 1, ordered=(nested.name == 'ol'))

        def add_table(table_element):
            rows = table_element.find_all('tr')
            if not rows:
                return
            cols = rows[0].find_all(['td', 'th'])
            table = doc.add_table(rows=1, cols=len(cols))
            table.style = 'Table Grid'
            for i, cell in enumerate(cols):
                table.rows[0].cells[i].text = cell.text.strip()
            for row in rows[1:]:
                cells = row.find_all(['td', 'th'])
                row_cells = table.add_row().cells
                for i, cell in enumerate(cells):
                    row_cells[i].text = cell.text.strip()
            doc.add_paragraph()

        for element in soup.contents:
            if not hasattr(element, 'name'):
                continue
            tag = element.name
            if tag in ['h1', 'h2', 'h3']:
                level = int(tag[1])
                p = doc.add_heading('', level=level)
                run = p.add_run(element.text.strip())
                run.bold = True
                p.paragraph_format.space_after = Pt(6)
                list_number = 1
            elif tag == 'p':
                p = doc.add_paragraph(element.text.strip())
                p.paragraph_format.space_after = Pt(6)
            elif tag == 'ul':
                process_list(element, indent_level=0, ordered=False)
            elif tag == 'ol':
                process_list(element, indent_level=0, ordered=True)
            elif tag == 'table':
                add_table(element)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as e:
        print(f"Error converting to DOCX: {e}")
        return None